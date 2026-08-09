"""Buduje treść OT.docx jako sekwencję: akapit z nazwą obrębu (zwykły
tekst), osobna tabela z opisem taksacyjnym tego obrębu, kolejny akapit,
kolejna tabela... Osobna tabela na obręb - nie da się tego zrobić samym
docxtpl/jinja `{%tr for %}` (ta pętla powtarza WIERSZE w jednej,
STAŁEJ tabeli w szablonie, a tu liczba samych TABEL jest zmienna) -
dlatego całość budowana jest tu, w Pythonie, przez python-docx,
i wstawiana do głównego szablonu jako docxtpl subdocument
({{p tabela_ot}} w szablony/szablon_ot.docx).

Wiersze nagłówka tabeli (5, scalenia, formatowanie) i wiersz danych
(formatowanie komórek/czcionki) są klonowane 1:1 z zasobu
szablony/szablon_ot_tabela_wzor.docx (wycięty z Twojego
SZABLON_OPERAT_UPUL.docx) - podmieniany jest tylko tekst w istniejących
runach (nie tworzone nowe), żeby nie zgubić formatowania. Scalanie
pionowe komórek adres/powierzchnia/opis (żeby nie powtarzały się w
każdym wierszu taksacji jednego wydzielenia) jest robione realnym
scaleniem komórek (cell.merge()), a nie sztuczką {% vm %} z docxtpl.
"""

import os
import re
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

_WZOR_SC = os.path.join(
    os.path.dirname(os.path.dirname(__file__)), 'szablony',
    'szablon_ot_tabela_wzor.docx')

# kolejnosc kolumn 0..17 tabeli OT (21 kolumn UPUL, ostatnie 3 to
# zawsze puste "Wykonane" - GeneratorOT ich nie liczy, tak samo jak w
# oryginale)
_KOLEJNOSC_POL = [
    'adr', 'pow', 'ops', 'war', 'gat', 'udz', 'wiek', 'zad', 'bhd',
    'wys', 'bon', 'vha', 'vpow', 'przyr', 'cue', 'cuep', 'cueb', 'cuen',
]

_wzor_cache = None


def _wczytaj_wzor():
    global _wzor_cache
    if _wzor_cache is not None:
        return _wzor_cache

    doc = Document(_WZOR_SC)
    tbl = doc.tables[0]._tbl
    wiersze = tbl.findall(qn('w:tr'))

    naglowek_wiersze = [deepcopy(w) for w in wiersze[:5]]
    wzor_wiersz = deepcopy(wiersze[5])
    tblPr = deepcopy(tbl.find(qn('w:tblPr')))
    tblGrid = deepcopy(tbl.find(qn('w:tblGrid')))

    _wzor_cache = (naglowek_wiersze, wzor_wiersz, tblPr, tblGrid)
    return _wzor_cache


def _ustaw_tekst(cell, wartosc):
    """Podmienia tekst PIERWSZEGO run w komórce (zachowując jego rPr -
    formatowanie sklonowane z wzorca), zamiast cell.text = ..., które
    tworzy nowy run i traci oryginalny font/rozmiar.

    run.text = ... w python-docx zamienia KAŻDY '\\n' na osobny
    <w:br/> - niektóre pola (np. przyrost w ot.py) mają w danych
    dosłowne ' \\n\\n ' jako separator (nieszkodliwe pod starym
    mechanizmem podstawiania tekstu przez docxtpl/jinja, które nie
    interpretowało '\\n' specjalnie) - tu dawałoby to dodatkową pustą
    linię, więc podwójne/wielokrotne złamania linii są ściągane do
    pojedynczego.
    """
    tekst = '' if wartosc in (None, '') else str(wartosc)
    tekst = re.sub(r'[ \t]*\n\s*\n[ \t]*', '\n', tekst)
    p = cell.paragraphs[0]
    if not p.runs:
        p.add_run(tekst)
        return
    p.runs[0].text = tekst
    for dodatkowy in p.runs[1:]:
        dodatkowy.text = ''


def _dodaj_tabele_obrebu(subdoc, tabela_ot, wzor):
    naglowek_wiersze, wzor_wiersz, tblPr, tblGrid = wzor
    ncol = len(naglowek_wiersze[-1].findall(qn('w:tc')))

    table = subdoc.add_table(rows=len(naglowek_wiersze), cols=ncol)
    tbl = table._tbl

    stare_tblPr = tbl.find(qn('w:tblPr'))
    tbl.replace(stare_tblPr, deepcopy(tblPr))
    stare_tblGrid = tbl.find(qn('w:tblGrid'))
    tbl.replace(stare_tblGrid, deepcopy(tblGrid))

    for stary, wzorcowy in zip(tbl.findall(qn('w:tr')), naglowek_wiersze):
        tbl.replace(stary, deepcopy(wzorcowy))

    for it in tabela_ot:
        opis = it.get('opis') or []
        if not opis:
            continue

        start_idx = len(table.rows)
        for wiersz in opis:
            tbl.append(deepcopy(wzor_wiersz))
            row = table.rows[-1]
            for i in range(ncol):
                pole = _KOLEJNOSC_POL[i] if i < len(_KOLEJNOSC_POL) else None
                wartosc = wiersz.get(pole, '') if pole else ''
                _ustaw_tekst(row.cells[i], wartosc)
        end_idx = len(table.rows) - 1

        if end_idx > start_idx:
            for col in (0, 1, 2):  # adr, pow, ops
                table.cell(start_idx, col).merge(table.cell(end_idx, col))

    return table


def zbuduj_subdoc(tpl, grupy):
    """grupy: [{'naglowek': 'Nazwa (kod)', 'tabela_ot': [...]}, ...] -
    zwraca docxtpl Subdoc gotowy do wstawienia jako {{p tabela_ot}}."""
    subdoc = tpl.new_subdoc()
    wzor = _wczytaj_wzor()

    for grupa in grupy:
        p = subdoc.add_paragraph()
        run = p.add_run(grupa['naglowek'])
        run.bold = True
        run.font.size = Pt(11)

        _dodaj_tabele_obrebu(subdoc, grupa['tabela_ot'], wzor)

    return subdoc
